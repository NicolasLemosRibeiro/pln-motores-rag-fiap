"""Gera o relatorio final unificado em DOCX com geometria deterministica."""

from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = ROOT / "outputs"
DOCX_PATH = DOCS / "relatorio_final_unificado.docx"
IMG_DIR = DOCS / "_report_assets"

BLUE = "2E74B5"
DARK = "0B2545"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
RED = "9B1C1C"
GOLD = "7A5A00"
WHITE = "FFFFFF"


def font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None: tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total)); tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None: tbl_pr.append(tbl_w)
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None: tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    if tbl_ind.getparent() is None: tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, header_fill=LIGHT):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(label)
        shade(cell, header_fill)
        for run in cell.paragraphs[0].runs: font(run, 9.5, True, DARK)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); tr_pr.append(repeat)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs: font(run, 9.2)
    table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p.add_run(text)
    return p


def add_callout(doc, label, text, color=DARK, fill=PALE_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single"); edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "5"); edge.set(qn("w:color"), "808080")
        borders.append(edge)
    p_pr.append(borders)
    r = p.add_run(label + " "); font(r, 10.5, True, color)
    r = p.add_run(text); font(r, 10.5, False, DARK)
    return p


def set_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, "1F4D78", 8, 4)
    ]:
        style = styles[name]; style.font.name = "Calibri"; style.font.size = Pt(size)
        style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"; style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8); style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("PLN INDUSTRIAL  |  SPRINTS 3 E 4")
    font(r, 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Prova de conceito sintética  |  Página "); font(r, 8.5, False, MUTED)
    set_page_field(footer)


def make_architecture(path):
    fig, ax = plt.subplots(figsize=(10, 3.8))
    ax.axis("off")
    boxes = [
        (0.02, 0.58, 0.18, 0.25, "Sensores e\nalertas", PALE_BLUE),
        (0.27, 0.58, 0.18, 0.25, "Narrativas +\nestado operacional", "DDEBF7"),
        (0.52, 0.58, 0.18, 0.25, "Prompt com\nmemória curta", "FFF2CC"),
        (0.77, 0.58, 0.18, 0.25, "Resposta, fonte\ne confiança", "E2F0D9"),
        (0.27, 0.08, 0.18, 0.25, "Documentos\ntécnicos", PALE_BLUE),
        (0.52, 0.08, 0.18, 0.25, "Chunks + índice\n+ re-ranking", "DDEBF7"),
    ]
    for x, y, w, h, label, fill in boxes:
        ax.add_patch(plt.Rectangle((x, y), w, h, facecolor="#" + fill, edgecolor="#" + DARK, linewidth=1.2))
        ax.text(x+w/2, y+h/2, label, ha="center", va="center", fontsize=10, color="#" + DARK, weight="bold")
    arrows = [((.20,.705),(.27,.705)),((.45,.705),(.52,.705)),((.70,.705),(.77,.705)),((.45,.205),(.52,.205)),((.61,.33),(.61,.58))]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", color="#"+BLUE, lw=2))
    ax.text(.50,.94,"Fluxo integrado da solução",ha="center",fontsize=14,weight="bold",color="#"+DARK)
    fig.tight_layout(); fig.savefig(path, dpi=180, bbox_inches="tight"); plt.close(fig)


def make_metrics(path, m3, m4_qwen):
    labels = ["F1 macro", "ROUGE-1", "Precision@1", "Hit@3", "Faithfulness"]
    values = [m3["classificacao"]["macro_f1"], m3["rouge"]["rouge1"], m4_qwen["precision_at_1"], m4_qwen["hit_at_3"], m4_qwen["faithfulness"]]
    fig, ax = plt.subplots(figsize=(8.8, 3.5))
    bars = ax.barh(labels, values, color=["#2E74B5", "#5B9BD5", "#70AD47", "#70AD47", "#ED7D31"])
    ax.set_xlim(0,1.08); ax.grid(axis="x", alpha=.2); ax.set_xlabel("Pontuação")
    for bar,value in zip(bars,values): ax.text(value+.015,bar.get_y()+bar.get_height()/2,f"{value:.3f}",va="center",fontsize=9)
    ax.set_title("Indicadores reproduzidos")
    fig.tight_layout(); fig.savefig(path,dpi=180,bbox_inches="tight"); plt.close(fig)


def build():
    DOCS.mkdir(exist_ok=True); IMG_DIR.mkdir(exist_ok=True)
    m3 = json.loads((OUT / "metricas_sprint3.json").read_text(encoding="utf-8"))
    m4 = json.loads((OUT / "metricas_sprint4.json").read_text(encoding="utf-8"))
    m4_qwen = json.loads((OUT / "metricas_qwen_hibrido.json").read_text(encoding="utf-8"))
    arch = IMG_DIR / "arquitetura.png"; metrics_img = IMG_DIR / "metricas.png"
    make_architecture(arch); make_metrics(metrics_img, m3, m4_qwen)
    doc = Document(); setup_styles(doc)

    # Capa editorial
    doc.add_paragraph().paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph(); kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("RELATÓRIO TÉCNICO UNIFICADO"); font(r, 10.5, True, GOLD)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Solução de PLN para Alertas\nOperacionais e Troubleshooting RAG"); font(r, 27, True, DARK)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; subtitle.paragraph_format.space_after = Pt(44)
    r = subtitle.add_run("Sprints 3 e 4  |  Motores elétricos  |  Prova de conceito reproduzível"); font(r, 13, False, BLUE)
    call = doc.add_paragraph(); call.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = call.add_run("30 alertas  •  200 eventos  •  24 chunks  •  20 perguntas"); font(r, 11, True, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(110)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Agosto de 2026\nDocumentos e dados sintéticos para demonstração acadêmica"); font(r, 10.5, False, MUTED)
    doc.add_page_break()

    doc.add_heading("Resumo executivo", level=1)
    doc.add_paragraph("A solução integra duas frentes complementares. A Sprint 3 transforma sinais numéricos em narrativas acessíveis, classifica eventos em cinco categorias e gera relatório operacional rastreável. A Sprint 4 usa documentação técnica e o estado real do alerta para recuperar evidências e responder perguntas de troubleshooting com fonte, confiança e memória curta.")
    add_callout(doc, "Resultado principal.", "O fluxo completo roda sem API paga. Além do baseline offline, a execução Colab usou Sentence Transformers, FAISS e Qwen 3B com guardrails e fallback rastreável.")
    picture = doc.add_picture(str(metrics_img), width=Inches(5.75))
    picture._inline.docPr.set("descr", "Gráfico horizontal com F1 macro, ROUGE-1, Precision@1, Hit@4 e faithfulness.")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figura 1. Indicadores obtidos na execução versionada. Valores em corpus sintético.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs: font(r, 9, False, MUTED, True)
    doc.add_heading("Escopo e premissas", level=2)
    add_bullet(doc, "Ativo demonstrado: motores elétricos MT Série 040, com foco no Motor MT-042.")
    add_bullet(doc, "Dados, limites e manuais são sintéticos e não substituem documentação do fabricante.")
    add_bullet(doc, "A solução apoia análise; não comanda o ativo, não libera trabalho e não substitui profissionais habilitados.")

    doc.add_heading("1. Conjunto de dados e rastreabilidade", level=1)
    doc.add_paragraph("A rastreabilidade foi tratada como requisito estrutural. Alertas preservam identificadores, equipamento, sensor, timestamp, baseline, valor atual, magnitude, unidade e janela temporal. Perguntas de troubleshooting apontam explicitamente para as seções documentais consideradas corretas.")
    add_table(doc, ["Conjunto", "Volume", "Finalidade"], [
        ("Alertas", "30", "Templates, ROUGE, checklist e relatório"),
        ("Eventos", "200", "Treino e holdout estratificado em cinco categorias"),
        ("Documentos", "3", "Manual, datasheet e ficha de manutenção"),
        ("Chunks", str(m4["chunks_indexados"]), "Índice vetorial e recuperação"),
        ("Perguntas", "20", "Gabarito de recuperação e resposta"),
    ], [2500, 1100, 5760])
    doc.add_heading("Campos de origem", level=2)
    doc.add_paragraph("Cada afirmação factual do relatório diário/semanal termina com uma referência no formato: [Fonte: alerta; sensor; valor; timestamp]. O assistente usa [documento > seção]. Essa separação permite auditar tanto a observação operacional quanto a orientação documental.")

    doc.add_heading("2. Sprint 3 — geração de linguagem natural", level=1)
    doc.add_paragraph("A geração é determinística e combina severidade, equipamento, tipo de sensor, desvio, direção em relação ao baseline, janela, horário e recomendação específica. O tom cresce com a urgência sem perder precisão.")
    add_table(doc, ["Nível", "Tom e vocabulário", "Ação típica"], [
        ("Leve", "Atenção e acompanhamento; evita alarmismo", "Monitorar tendência e conferir na ronda"),
        ("Moderado", "Alerta explícito; destaca desvio e janela", "Inspeção direcionada antes de ampliar carga"),
        ("Crítico", "Risco e indisponibilidade em linguagem direta", "Parada segura, isolamento ou não reenergização"),
    ], [1300, 4050, 4010])
    doc.add_heading("Exemplos dos três templates", level=2)
    examples = [
        ("Leve", "Atenção: desvio leve no Motor MT-042. A temperatura do enrolamento variou +8°C acima do baseline. Acompanhar a tendência. [Fonte: alerta ALT-0001; sensor TMP-042-1]"),
        ("Moderado", "Alerta moderado detectado no Motor MT-017. A vibração do mancal apresentou desvio de +2,3 mm/s RMS. Recomenda-se inspecionar alinhamento, fixações e mancais. [Fonte: alerta ALT-0002; sensor VIB-043-2]"),
        ("Crítico", "ALERTA CRÍTICO no Motor MT-105: a corrente de fase atingiu +62 A acima do baseline. Há risco de dano ou indisponibilidade. Isolar conforme procedimento e inspecionar alimentação e conexões. [Fonte: alerta ALT-0003; sensor CUR-044-1]"),
    ]
    for label, text in examples: add_callout(doc, label + ".", text, RED if label == "Crítico" else DARK, "FDECEC" if label == "Crítico" else PALE_BLUE)

    doc.add_heading("3. Classificação textual e relatório", level=1)
    doc.add_paragraph("O classificador usa TF-IDF de unigramas e bigramas, seguido de regressão logística com balanceamento. A avaliação emprega holdout estratificado de 25% e semente fixa. As cinco categorias e seus critérios são apresentados abaixo.")
    add_table(doc, ["Categoria", "Critério linguístico-operacional"], [
        ("Manutenção corretiva", "Reparo ou substituição executada após falha"),
        ("Manutenção preventiva", "Inspeção, limpeza, lubrificação ou ensaio programado"),
        ("Anomalia elétrica", "Corrente, tensão, fase, isolamento, borne ou proteção"),
        ("Anomalia mecânica", "Vibração, ruído, mancal, alinhamento, folga ou lubrificação"),
        ("Operação normal", "Estabilidade, disponibilidade e ausência de desvios"),
    ], [2850, 6510])
    doc.add_heading("Métricas", level=2)
    rows = [(k, f"{v:.3f}") for k, v in m3["classificacao"]["f1_por_categoria"].items()]
    add_table(doc, ["Categoria", "F1"], rows + [("Macro F1", f"{m3['classificacao']['macro_f1']:.3f}")], [7600, 1760])
    add_callout(doc, "Leitura responsável.", "O F1 de 1,000 é coerente com frases sintéticas balanceadas e regulares. O relatório não extrapola esse valor para dados reais.", GOLD, "FFF7E0")
    add_table(doc, ["Métrica de geração", "Valor"], [
        ("ROUGE-1", f"{m3['rouge']['rouge1']:.3f}"), ("ROUGE-2", f"{m3['rouge']['rouge2']:.3f}"),
        ("ROUGE-L", f"{m3['rouge']['rougeL']:.3f}"), ("Clareza", "100%"),
        ("Precisão de identificadores", "100%"), ("Utilidade lexical", "80%"), ("Rastreabilidade", "100%"),
    ], [7600, 1760])
    doc.add_heading("Relatório operacional", level=2)
    report = (OUT / "relatorio_operacional.txt").read_text(encoding="utf-8")
    for paragraph in report.split("\n\n"):
        doc.add_paragraph(paragraph)

    doc.add_heading("4. Sprint 4 — arquitetura RAG", level=1)
    picture = doc.add_picture(str(arch), width=Inches(5.85))
    picture._inline.docPr.set("descr", "Arquitetura que integra sensores, narrativas, documentos, chunks, prompt e resposta fundamentada.")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figura 2. Contexto operacional e documentação convergem no prompt do assistente.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs: font(r, 9, False, MUTED, True)
    doc.add_heading("Chunking e metadados", level=2)
    doc.add_paragraph("O chunking preserva cabeçalhos Markdown e mantém parágrafos inteiros. Se uma seção excede 1.200 caracteres, a divisão ocorre na fronteira entre parágrafos com sobreposição de 160 caracteres. Os metadados registram documento, seção, equipamento e tipos de anomalia.")
    doc.add_heading("Embeddings e índice", level=2)
    doc.add_paragraph(f"A execução de referência usou {m4['backend']}: vetores TF-IDF e índice vetorial NearestNeighbors por similaridade de cosseno. O modo opcional usa embeddings multilíngues Sentence Transformers e FAISS; na ausência de FAISS, mantém índice equivalente por NearestNeighbors.")
    doc.add_heading("Retriever e re-ranking", level=2)
    add_table(doc, ["Componente do score", "Peso/base", "Objetivo"], [
        ("Similaridade do conteúdo", "50%", "Encontrar trechos com termos e relações da pergunta"),
        ("Similaridade do cabeçalho", "30%", "Preservar intenção técnica da seção"),
        ("Cobertura lexical", "20%", "Dar estabilidade a termos de domínio"),
        ("Bônus de metadados", "Reduzido", "Priorizar anomalia, equipamento e severidade"),
    ], [3200, 1500, 4660])

    doc.add_heading("5. Assistente conversacional", level=1)
    doc.add_paragraph("O prompt define persona de especialista em motores, exige resposta fundamentada, fonte, confiança, linguagem segura e abstenção quando faltar evidência. O estado do equipamento é injetado como bloco separado e inclui resumo do alerta, sensor, severidade e tipo de anomalia.")
    add_callout(doc, "Regra de segurança.", "O modelo não deve instruir intervenção energizada, inventar valores ausentes nem apresentar diagnóstico conclusivo com evidência insuficiente.", RED, "FDECEC")
    doc.add_heading("Memória curta", level=2)
    doc.add_paragraph("Uma fila limitada aos quatro últimos turnos mantém referência a perguntas anteriores sem carregar indefinidamente informações possivelmente obsoletas. Cada nova pergunta aciona recuperação documental novamente.")
    doc.add_heading("Backends de geração", level=2)
    add_table(doc, ["Backend", "Uso", "Trade-off"], [
        ("Extrativo", "Execução offline padrão", "Mais auditável; menor naturalidade"),
        ("Qwen 3B protegido", "Executado em Colab T4", "Mais natural; exige validação e fallback"),
        ("API compatível", "Ambiente com chave e governança", "Escalável; custo, privacidade e dependência externa"),
    ], [2100, 3300, 3960])
    add_callout(doc, "Guardrail híbrido.", "Cada saída do Qwen é validada quanto a citações, sustentação, abstenção indevida, sequência de procedimento e contradição de segurança. Respostas reprovadas são revisadas uma vez e, se necessário, substituídas por fallback extrativo.", RED, "FDECEC")

    doc.add_heading("6. Avaliação da solução conversacional", level=1)
    doc.add_paragraph("O conjunto de 20 perguntas foi construído diretamente das seções sintéticas. Cada item contém pergunta, resposta de referência, seções relevantes e contexto operacional. A tabela separa qualidade de recuperação e qualidade da resposta.")
    add_table(doc, ["Métrica", "Resultado", "Interpretação"], [
        ("Precision@1", f"{m4_qwen['precision_at_1']:.3f}", "Seção correta na primeira posição"),
        ("Hit@3", f"{m4_qwen['hit_at_3']:.3f}", "Ao menos uma seção correta entre três"),
        ("MRR", f"{m4_qwen['mrr']:.3f}", "Posição média do primeiro acerto"),
        ("Context precision@3", f"{m4_qwen['context_precision']:.3f}", "Chunks relevantes entre os três"),
        ("Faithfulness", f"{m4_qwen['faithfulness']:.3f}", "Suporte sentencial no documento/contexto"),
        ("Answer relevancy", f"{m4_qwen['answer_relevancy']:.3f}", "Cobertura do gabarito + cosseno"),
    ], [2600, 1500, 5260])
    modes = m4_qwen["generation_modes"]
    add_callout(doc, "Resultado dos guardrails.", f"Das 20 respostas: {modes.get('llm', 0)} foram aceitas diretamente, {modes.get('llm_revised', 0)} após revisão automática e {modes.get('guardrail_fallback', 0)} acionaram fallback rastreável.", GOLD, "FFF7E0")
    add_callout(doc, "Nota metodológica.", "As métricas de resposta são proxies locais transparentes. Uma avaliação de produção deve combinar RAGAS ou juiz independente com revisão de especialistas e testes adversariais.", GOLD, "FFF7E0")

    doc.add_heading("7. Três cenários de demonstração", level=1)
    demo = json.loads((OUT / "demonstracao_qwen_hibrido.json").read_text(encoding="utf-8"))
    scenarios = demo["cenarios"]
    for item in scenarios:
        doc.add_heading(item["cenario"].capitalize(), level=2)
        p = doc.add_paragraph(); r = p.add_run("Contexto: "); font(r, 11, True, DARK); p.add_run(item["contexto"]["resumo_alerta"])
        p = doc.add_paragraph(); r = p.add_run("Resposta demonstrada: "); font(r, 11, True, DARK); p.add_run(item["answer"])
        p = doc.add_paragraph(); r = p.add_run("Fontes recuperadas: "); font(r, 10, True, MUTED); p.add_run("; ".join(item["sources"][:2]))
        p = doc.add_paragraph(); r = p.add_run("Modo: "); font(r, 10, True, MUTED); p.add_run(item["generation_mode"])
    memory = demo["demonstracao_memoria"]
    add_callout(doc, "Memória demonstrada.", f"A pergunta de continuação recebeu o turno anterior no prompt; memória injetada = {memory['memoria_anterior_injetada']} e turnos armazenados = {memory['turnos_armazenados']}.")

    doc.add_heading("8. Limites, riscos e mitigação", level=1)
    doc.add_heading("Fora do escopo", level=2)
    for text in [
        "Outros ativos, marcas, tensões ou componentes ausentes do corpus.",
        "Torque, tolerância ou limite não explicitamente documentado.",
        "Diagnóstico conclusivo a partir de um único sensor.",
        "Instruções para contornar bloqueio, proteção ou intertravamento.",
        "Decisões autônomas de parada, partida ou retorno ao serviço.",
    ]: add_bullet(doc, text)
    doc.add_heading("Alucinações e mitigação", level=2)
    add_table(doc, ["Risco", "Mitigação implementada"], [
        ("Combinar procedimentos em ordem incorreta", "Chunking por seção, fontes e prompt restritivo"),
        ("Confundir sensores semelhantes", "Metadados de sensor e contexto operacional separado"),
        ("Inventar valor ou estado ausente", "Validação pós-geração e fallback extrativo"),
        ("Negar critério de parada segura", "Regra determinística confronta resposta e evidência"),
        ("Carregar contexto obsoleto", "Memória limitada a quatro turnos"),
        ("Métrica automática otimista", "Limitações explícitas e recomendação de revisão humana"),
    ], [3900, 5460])

    doc.add_heading("9. Reprodução, entrega e próximos passos", level=1)
    doc.add_paragraph("Os notebooks foram executados integralmente e preservam as saídas. Scripts equivalentes permitem gerar dados, métricas e cenários pela linha de comando. O repositório inclui testes, documentação e artefatos JSON da execução Qwen protegida.")
    add_table(doc, ["Entregável", "Arquivo", "Estado"], [
        ("Notebook Sprint 3", "sprint3_pln_alertas.ipynb", "Executado sem erros"),
        ("Notebook Sprint 4", "sprint4_pln_rag.ipynb", "Executado sem erros"),
        ("Dados e gabaritos", "data/", "30 alertas, 200 eventos, 20 perguntas"),
        ("Relatório final", "docs/relatorio_final_unificado.docx", "Atualizado com Qwen 3B"),
        ("Execução Qwen", "outputs/", "Métricas, 20 perguntas e três cenários preservados"),
    ], [2800, 4300, 2260])
    doc.add_heading("Critérios antes de produção", level=2)
    for text in [
        "Substituir documentos sintéticos por corpus controlado, versionado e aprovado.",
        "Rotular registros reais e reavaliar F1, viés e drift por categoria.",
        "Validar 20 perguntas e novos casos adversariais com engenharia e segurança.",
        "Calibrar limiar de abstenção, controle de acesso, logs e retenção de dados.",
        "Realizar piloto sem comando do ativo e com aprovação humana obrigatória.",
    ]: add_bullet(doc, text)

    doc.add_heading("Conclusão", level=1)
    doc.add_paragraph("A solução cumpre o fluxo técnico solicitado: transforma alertas em linguagem natural, classifica eventos, produz relatório rastreável, indexa documentos, recupera contexto relevante e responde com segurança, fonte, confiança e memória curta. Os resultados demonstram viabilidade funcional; as ressalvas deixam claro o trabalho necessário para validação em ambiente real.")

    doc.core_properties.title = "Solução de PLN para Alertas Operacionais e Troubleshooting RAG"
    doc.core_properties.subject = "Relatório técnico unificado das Sprints 3 e 4"
    doc.core_properties.author = "Projeto PLN Industrial"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
